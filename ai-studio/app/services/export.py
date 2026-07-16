import os
import re
import hashlib
import logging
import asyncio
from concurrent.futures import ThreadPoolExecutor
from html.parser import HTMLParser

import markdown
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

from app.config import settings

logger = logging.getLogger(__name__)

MATH_BLOCK_RE = re.compile(r'\$\$(.*?)\$\$', re.DOTALL)
JOB_ID_RE = re.compile(r'^[a-zA-Z0-9_-]+$')


def _get_thread_pool() -> ThreadPoolExecutor:
    return ThreadPoolExecutor(max_workers=4)


def _wrap_math_for_html(rendered_html: str) -> str:
    def replace_math(m: re.Match[str]) -> str:
        math_content = m.group(1).strip()
        escaped = (
            math_content
            .replace('&', '&amp;')
            .replace('<', '&lt;')
            .replace('>', '&gt;')
        )
        return (
            '<div class="math-block" style="'
            'font-family: Courier New, Courier, monospace; '
            'background-color: #f8f9fa; '
            'border: 1px solid #dee2e6; '
            'border-radius: 4px; '
            'padding: 8px 12px; '
            'margin: 8px 0; '
            'direction: ltr; '
            'text-align: left; '
            'font-size: 10pt; '
            'white-space: pre-wrap; '
            'overflow-wrap: break-word;'
            '">' + escaped + '</div>'
        )
    return MATH_BLOCK_RE.sub(replace_math, rendered_html)


def save_markdown(content: str, output_path: str) -> str:
    parent_dir = os.path.dirname(output_path)
    if parent_dir:
        os.makedirs(parent_dir, exist_ok=True)
    try:
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(content)
    except OSError as exc:
        raise RuntimeError(f"Markdown save failed: {exc}") from exc
    return output_path


def _sanitize_job_id(job_id: str) -> str:
    if not JOB_ID_RE.match(job_id):
        raise ValueError(f"Invalid job_id: {job_id}")
    return job_id


def _register_fonts(pdf: FPDF) -> None:
    font_candidates = [
        ("C:\\Windows\\Fonts\\segoeui.ttf", "C:\\Windows\\Fonts\\segoeuib.ttf"),
        ("C:\\Windows\\Fonts\\arial.ttf", "C:\\Windows\\Fonts\\arialbd.ttf"),
        ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
        ("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf", "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf"),
        ("/usr/share/fonts/TTF/DejaVuSans.ttf", "/usr/share/fonts/TTF/DejaVuSans-Bold.ttf"),
        ("/System/Library/Fonts/Helvetica.ttc", "/System/Library/Fonts/Helvetica-Bold.ttc"),
        ("/System/Library/Fonts/Arial.ttf", "/System/Library/Fonts/Arial-Bold.ttf"),
        ("/Library/Fonts/Arial.ttf", "/Library/Fonts/Arial-Bold.ttf"),
    ]
    registered = False
    for regular_path, bold_path in font_candidates:
        if os.path.exists(regular_path):
            try:
                pdf.add_font("body", "", regular_path)
                if os.path.exists(bold_path):
                    pdf.add_font("body", "B", bold_path)
                pdf.set_font("body", "", 11)
                registered = True
                break
            except Exception:
                continue
    if not registered:
        pdf.set_font("Helvetica", "", 11)
        pdf.set_font("Helvetica", "B", 11)


def _render_markdown_to_pdf_sync(pdf: FPDF, md_content: str) -> None:
    available_width = pdf.w - pdf.l_margin - pdf.r_margin
    if available_width <= 0:
        raise ValueError(
            f"PDF available width is {available_width:.1f}pt; "
            f"page width={pdf.w:.1f}pt, margins={pdf.l_margin:.1f}/{pdf.r_margin:.1f}pt"
        )

    in_code_block = False
    code_buffer: list[str] = []

    lines = md_content.split("\n")
    for line in lines:
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code_block:
                pdf.set_font("Courier", "", 10)
                pdf.set_fill_color(248, 249, 250)
                code_text = "\n".join(code_buffer)
                if code_text:
                    pdf.multi_cell(available_width, 6, code_text, fill=True)
                pdf.ln(1)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_buffer.append(stripped)
            continue

        if stripped.startswith("# "):
            pdf.set_font("body", "B", 16)
            text = stripped[2:]
            if text:
                pdf.multi_cell(available_width, 8, text)
            pdf.ln(2)
        elif stripped.startswith("## "):
            pdf.set_font("body", "B", 14)
            text = stripped[3:]
            if text:
                pdf.multi_cell(available_width, 7, text)
            pdf.ln(1)
        elif stripped.startswith("### "):
            pdf.set_font("body", "B", 12)
            text = stripped[4:]
            if text:
                pdf.multi_cell(available_width, 6, text)
            pdf.ln(1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            pdf.set_font("body", "", 11)
            text = stripped[2:]
            if text:
                pdf.multi_cell(available_width, 6, f"\u2022 {text}")
        elif stripped == "---":
            pdf.ln(4)
        elif stripped == "":
            pdf.ln(3)
        else:
            if stripped.startswith("$$") and stripped.endswith("$$"):
                pdf.set_font("Courier", "", 10)
                text = stripped[2:-2]
                if text:
                    pdf.multi_cell(available_width, 6, text)
                pdf.ln(1)
            else:
                pdf.set_font("body", "", 11)
                if stripped:
                    pdf.multi_cell(available_width, 6, stripped)


def _render_docx_from_ast(doc: Document, md_content: str) -> None:
    try:
        html = markdown.markdown(md_content, extensions=["tables", "fenced_code", "codehilite"])

        class DocxHTMLParser(HTMLParser):
            def __init__(self, doc: Document) -> None:
                super().__init__()
                self.doc = doc
                self.in_table = False
                self.table_rows: list[list[str]] = []
                self.current_row: list[str] = []

            def handle_starttag(self, tag: str, attrs: list) -> None:
                if tag == "table":
                    self.in_table = True
                    self.table_rows = []
                elif tag == "tr" and self.in_table:
                    self.current_row = []
                elif tag in ("td", "th") and self.in_table:
                    self.current_row.append("")

            def handle_endtag(self, tag: str) -> None:
                if tag == "tr" and self.in_table:
                    self.table_rows.append(self.current_row)
                    self.current_row = []
                elif tag == "table":
                    self._render_table()
                    self.in_table = False

            def handle_data(self, data: str) -> None:
                stripped = data.strip()
                if not stripped:
                    return
                if "$$" in stripped:
                    parts = stripped.split("$$")
                    for i, part in enumerate(parts):
                        if i % 2 == 1:
                            self.doc.add_paragraph(f"[equation: {part.strip()}]")
                        elif part:
                            self.doc.add_paragraph(part)
                else:
                    self.doc.add_paragraph(stripped)

            def _render_table(self) -> None:
                if not self.table_rows:
                    return
                max_cols = max(len(row) for row in self.table_rows)
                table = self.doc.add_table(rows=len(self.table_rows), cols=max_cols)
                table.style = "Table Grid"
                for row_idx, row in enumerate(self.table_rows):
                    for col_idx, cell_text in enumerate(row):
                        if col_idx < max_cols:
                            table.rows[row_idx].cells[col_idx].text = cell_text
                self.table_rows = []

        parser = DocxHTMLParser(doc)
        parser.feed(html)
    except Exception as render_exc:
        logger.error("DOCX AST rendering failed, using line-by-line fallback: %s", render_exc, exc_info=True)
        lines = md_content.split("\n")
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            if stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("- "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            elif stripped.startswith("$$") and stripped.endswith("$$"):
                doc.add_paragraph(f"[equation: {stripped[2:-2]}]")
            elif MATH_BLOCK_RE.search(stripped):
                doc.add_paragraph(f"[equation: {stripped}]")
            else:
                doc.add_paragraph(stripped)


def _render_pdf_sync(markdown_path: str, output_path: str) -> None:
    try:
        with open(markdown_path, "r", encoding="utf-8") as f:
            md_content = f.read()
    except OSError as exc:
        raise RuntimeError(f"PDF export failed: {exc}") from exc

    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_margins(15, 15, 15)

        _register_fonts(pdf)
        _render_markdown_to_pdf_sync(pdf, md_content)
        pdf.output(output_path)
    except Exception as exc:
        raise RuntimeError(f"PDF export failed: {exc}") from exc


def _render_docx_sync(markdown_path: str, output_path: str) -> None:
    try:
        with open(markdown_path, "r", encoding="utf-8") as f:
            md_content = f.read()
    except OSError as exc:
        raise RuntimeError(f"DOCX export failed: {exc}") from exc

    try:
        doc = Document()
        _render_docx_from_ast(doc, md_content)
        doc.save(output_path)
    except Exception as exc:
        raise RuntimeError(f"DOCX export failed: {exc}") from exc


def _render_txt_sync(markdown_path: str, output_path: str) -> None:
    try:
        with open(markdown_path, "r", encoding="utf-8") as f:
            md_content = f.read()
    except OSError as exc:
        raise RuntimeError(f"TXT export failed: {exc}") from exc

    try:
        text = re.sub(r'#+\s+', '', md_content)
        text = re.sub(r'\*{1,2}(.+?)\*{1,2}', r'\1', text)
        text = re.sub(r'`{1,3}[^`]*`{1,3}', '', text)
        text = re.sub(r'!\[.*?\]\(.*?\)', '', text)
        text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
        text = re.sub(r'^\s*[-*+]\s+', '', text, flags=re.MULTILINE)
        text = re.sub(r'^\s*\d+\.\s+', '', text, flags=re.MULTILINE)
        text = re.sub(r'^\s*\|.+\|\s*$', '', text, flags=re.MULTILINE)
        text = re.sub(r'\$\$.*?\$\$', '', text, flags=re.DOTALL)
        text = re.sub(r'\$.*?\$', '', text)
        text = re.sub(r'^-{3,}$', '', text, flags=re.MULTILINE)
        text = re.sub(r'\n{3,}', '\n\n', text)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(text.strip())
    except Exception as exc:
        raise RuntimeError(f"TXT export failed: {exc}") from exc


def _render_html_sync(markdown_path: str, output_path: str) -> None:
    try:
        with open(markdown_path, "r", encoding="utf-8") as f:
            md_content = f.read()
    except OSError as exc:
        raise RuntimeError(f"HTML export failed: {exc}") from exc

    try:
        html_content = markdown.markdown(
            md_content,
            extensions=["tables", "fenced_code", "codehilite", "nl2br"],
        )
        html_content = _wrap_math_for_html(html_content)
        full_html = (
            "<!DOCTYPE html>"
            '<html lang="ar" dir="rtl">'
            "<head>"
            '<meta charset="utf-8">'
            '<meta name="viewport" content="width=device-width, initial-scale=1">'
            "<title>AI Studio Result</title>"
            "<style>"
            "body { font-family: 'Tajawal', 'Cairo', sans-serif; padding: 40px; max-width: 900px; margin: 0 auto; direction: rtl; text-align: right; line-height: 1.8; color: #1f2937; }"
            "h1, h2, h3 { margin-top: 1.5em; margin-bottom: 0.5em; }"
            "table { border-collapse: collapse; width: 100%; margin: 1em 0; }"
            "th, td { border: 1px solid #e5e7eb; padding: 8px 12px; text-align: right; }"
            "th { background: #f9fafb; font-weight: 600; }"
            "pre { background: #f8f9fa; padding: 12px; border-radius: 6px; overflow-x: auto; direction: ltr; text-align: left; }"
            "code { font-family: 'IBM Plex Mono', monospace; font-size: 0.9em; background: #f3f4f6; padding: 2px 6px; border-radius: 4px; }"
            "pre code { background: transparent; padding: 0; }"
            ".math-block { font-family: 'IBM Plex Mono', 'Fira Code', monospace; background: #f8f9fa; border: 1px solid #dee2e6; border-radius: 4px; padding: 8px 12px; margin: 8px 0; direction: ltr; text-align: left; font-size: 10pt; white-space: pre-wrap; overflow-wrap: break-word; }"
            "blockquote { border-right: 4px solid #e5e7eb; padding-right: 16px; margin: 1em 0; color: #6b7280; }"
            "hr { border: none; border-top: 1px solid #e5e7eb; margin: 2em 0; }"
            "</style>"
            "</head>"
            "<body>"
            + html_content
            + "</body></html>"
        )
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(full_html)
    except Exception as exc:
        raise RuntimeError(f"HTML export failed: {exc}") from exc


def get_thread_pool() -> ThreadPoolExecutor:
    return _get_thread_pool()


async def _export_async(
    markdown_path: str,
    output_pdf: str,
    output_docx: str,
    output_txt: str = "",
    output_html: str = "",
) -> dict[str, list[str]]:
    loop = asyncio.get_running_loop()
    thread_pool = _get_thread_pool()
    export_warnings: list[str] = []

    tasks = [
        loop.run_in_executor(thread_pool, _render_pdf_sync, markdown_path, output_pdf),
        loop.run_in_executor(thread_pool, _render_docx_sync, markdown_path, output_docx),
    ]
    if output_txt:
        tasks.append(loop.run_in_executor(thread_pool, _render_txt_sync, markdown_path, output_txt))
    if output_html:
        tasks.append(loop.run_in_executor(thread_pool, _render_html_sync, markdown_path, output_html))

    results = await asyncio.gather(*tasks, return_exceptions=True)
    for result in results:
        if isinstance(result, Exception):
            export_warnings.append(f"Export failed: {type(result).__name__}: {result}")
            logger.error("Export task failed: %s", result, exc_info=result)

    return {"export_warnings": export_warnings}


def markdown_to_pdf(markdown_path: str, output_path: str) -> str:
    _render_pdf_sync(markdown_path, output_path)
    return output_path


def markdown_to_docx(markdown_path: str, output_path: str) -> str:
    _render_docx_sync(markdown_path, output_path)
    return output_path


def markdown_to_txt(markdown_path: str, output_path: str) -> str:
    _render_txt_sync(markdown_path, output_path)
    return output_path


def markdown_to_html(markdown_path: str, output_path: str) -> str:
    _render_html_sync(markdown_path, output_path)
    return output_path


async def export_markdown_async(
    markdown_path: str,
    output_pdf: str,
    output_docx: str,
    output_txt: str = "",
    output_html: str = "",
) -> dict[str, list[str]]:
    return await _export_async(markdown_path, output_pdf, output_docx, output_txt, output_html)


def re_export_job(
    job_id: str,
    markdown: str,
    output_pdf: str = "",
    output_docx: str = "",
    output_txt: str = "",
    output_html: str = "",
) -> dict:
    _sanitize_job_id(job_id)
    job_dir = os.path.join(settings.upload_dir, job_id)
    if not os.path.isdir(job_dir):
        raise RuntimeError(f"Job directory not found: {job_dir}")

    md_path = os.path.join(job_dir, "result.md")
    save_markdown(markdown, md_path)

    export_warnings: list[str] = []
    pdf_path = output_pdf or os.path.join(job_dir, "result.pdf")
    docx_path = output_docx or os.path.join(job_dir, "result.docx")
    txt_path = output_txt or os.path.join(job_dir, "result.txt")
    html_path = output_html or os.path.join(job_dir, "result.html")

    futures = []
    if pdf_path:
        futures.append(_get_thread_pool().submit(_render_pdf_sync, md_path, pdf_path))
    if docx_path:
        futures.append(_get_thread_pool().submit(_render_docx_sync, md_path, docx_path))
    if txt_path:
        futures.append(_get_thread_pool().submit(_render_txt_sync, md_path, txt_path))
    if html_path:
        futures.append(_get_thread_pool().submit(_render_html_sync, md_path, html_path))

    for future in futures:
        try:
            future.result()
        except Exception as exc:
            export_warnings.append(f"Export failed: {type(exc).__name__}: {exc}")

    base_url = f"/uploads/{job_id}"
    return {
        "markdown_path": f"{base_url}/result.md",
        "pdf_path": f"{base_url}/result.pdf" if pdf_path and os.path.exists(pdf_path) else None,
        "docx_path": f"{base_url}/result.docx" if docx_path and os.path.exists(docx_path) else None,
        "txt_path": f"{base_url}/result.txt" if txt_path and os.path.exists(txt_path) else None,
        "html_path": f"{base_url}/result.html" if html_path and os.path.exists(html_path) else None,
        "export_warnings": export_warnings,
    }
